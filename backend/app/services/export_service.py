from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Optional
from datetime import datetime

class ExportService:
    @staticmethod
    def export_to_docx(planificacion: dict, content: str) -> BytesIO:
        """
        Genera un documento Word a partir de la planificación.
        """
        doc = Document()

        # Encabezado
        title = doc.add_heading(
            planificacion.get("titulo", "Planificación sin título"),
            level=1
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadatos
        metadata = doc.add_paragraph()
        metadata.add_run("Tipo de Material: ").bold = True
        metadata.add_run(planificacion.get("tipo_material", "N/A"))

        metadata2 = doc.add_paragraph()
        metadata2.add_run("Fecha de Creación: ").bold = True
        metadata2.add_run(
            datetime.fromisoformat(planificacion.get("created_at", datetime.utcnow().isoformat())).strftime("%d/%m/%Y")
        )

        # Idea Inicial
        doc.add_heading("Idea Inicial", level=2)
        doc.add_paragraph(planificacion.get("idea_inicial", ""))

        # Contenido Principal
        doc.add_heading("Contenido", level=2)
        doc.add_paragraph(content)

        # Pie de página
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Generado con PlanificaTool - Herramienta de Planificación Educativa con IA"

        # Guardar en BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def export_to_pdf_simple(planificacion: dict, content: str) -> BytesIO:
        """
        Genera un PDF simple a partir de la planificación.
        Esta versión usa reportlab para crear un PDF básico.
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors

        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()

        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#134E5E'),
            spaceAfter=30,
            alignment=1  # Center
        )
        title = Paragraph(planificacion.get("titulo", "Planificación sin título"), title_style)
        story.append(title)

        # Metadata
        meta_data = [
            ["Tipo de Material:", planificacion.get("tipo_material", "N/A")],
            ["Fecha de Creación:", datetime.fromisoformat(planificacion.get("created_at", datetime.utcnow().isoformat())).strftime("%d/%m/%Y")]
        ]
        meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F0F0F0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 0.3*inch))

        # Idea Inicial
        story.append(Paragraph("Idea Inicial", styles['Heading2']))
        story.append(Paragraph(planificacion.get("idea_inicial", ""), styles['BodyText']))
        story.append(Spacer(1, 0.2*inch))

        # Contenido
        story.append(Paragraph("Contenido", styles['Heading2']))
        story.append(Paragraph(content, styles['BodyText']))

        # Construir PDF
        doc.build(story)
        output.seek(0)
        return output

    @staticmethod
    def export_to_pdf(planificacion: dict, content: str) -> BytesIO:
        """
        Exporta a PDF. Por ahora usa la versión simple.
        Puede mejorarse con html2pdf o similar.
        """
        return ExportService.export_to_pdf_simple(planificacion, content)
